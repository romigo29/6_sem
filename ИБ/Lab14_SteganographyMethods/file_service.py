# file_service.py
# Работа с файлами-контейнерами text.txt и text.docx.

import os
import textwrap
from docx import Document

from config import INPUT_TXT, INPUT_DOCX, AUTO_LINE_WIDTH


def get_app_dir() -> str:
    """Возвращает папку, где расположен main.py."""
    return os.path.dirname(os.path.abspath(__file__))


def get_path(filename: str) -> str:
    """Формирует путь к файлу в одной директории с main.py."""
    return os.path.join(get_app_dir(), filename)


def load_container_document(for_line_spacing: bool = False) -> Document:
    """
    Загружает контейнер.
    Приоритет: text.docx, затем text.txt.
    Если найден text.txt, из него создается объект Document.
    """
    docx_path = get_path(INPUT_DOCX)
    txt_path = get_path(INPUT_TXT)

    if os.path.exists(docx_path):
        return Document(docx_path)

    if os.path.exists(txt_path):
        with open(txt_path, "r", encoding="utf-8") as file:
            text = file.read()
        return create_document_from_text(text, for_line_spacing=for_line_spacing)

    raise FileNotFoundError(
        f"Не найден файл-контейнер. Поместите {INPUT_TXT} или {INPUT_DOCX} в папку с main.py."
    )


def create_document_from_text(text: str, for_line_spacing: bool = False) -> Document:
    """
    Создает DOCX из обычного текста.
    Для метода межстрочного интервала текст автоматически делится на короткие строки,
    чтобы увеличить количество абзацев-носителей.
    """
    doc = Document()
    raw_lines = [line.strip() for line in text.splitlines() if line.strip()]

    if not raw_lines:
        raw_lines = [text.strip()] if text.strip() else ["Пустой контейнер."]

    if for_line_spacing:
        paragraphs = []
        for line in raw_lines:
            paragraphs.extend(textwrap.wrap(line, width=AUTO_LINE_WIDTH) or [line])
    else:
        paragraphs = raw_lines

    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)

    return doc


def save_document(doc: Document, filename: str) -> str:
    """Сохраняет DOCX в папке с main.py."""
    path = get_path(filename)
    doc.save(path)
    return path

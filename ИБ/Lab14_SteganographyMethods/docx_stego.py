# docx_stego.py
# Методы внедрения и извлечения скрытого сообщения из DOCX.

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from binary_codec import text_to_bits, bits_to_text
from config import (
    LINE_ZERO_TWIPS,
    LINE_ONE_TWIPS,
    LINE_THRESHOLD_TWIPS,
    KERNING_ZERO_TWIPS,
    KERNING_ONE_TWIPS,
    KERNING_THRESHOLD_TWIPS,
)


def _set_paragraph_line_spacing(paragraph, twips: int) -> None:
    """Записывает точный межстрочный интервал в XML абзаца."""
    p_pr = paragraph._p.get_or_add_pPr()
    spacing = p_pr.find(qn("w:spacing"))

    if spacing is None:
        spacing = OxmlElement("w:spacing")
        p_pr.append(spacing)

    spacing.set(qn("w:line"), str(twips))
    spacing.set(qn("w:lineRule"), "exact")


def _get_paragraph_line_spacing(paragraph) -> int | None:
    """Читает межстрочный интервал из XML абзаца."""
    p_pr = paragraph._p.pPr
    if p_pr is None:
        return None

    spacing = p_pr.find(qn("w:spacing"))
    if spacing is None:
        return None

    value = spacing.get(qn("w:line"))
    return int(value) if value is not None else None


def _set_run_spacing(run, twips: int) -> None:
    """Записывает межсимвольный интервал run в XML."""
    r_pr = run._r.get_or_add_rPr()
    spacing = r_pr.find(qn("w:spacing"))

    if spacing is None:
        spacing = OxmlElement("w:spacing")
        r_pr.append(spacing)

    spacing.set(qn("w:val"), str(twips))


def _get_run_spacing(run) -> int | None:
    """Читает межсимвольный интервал run из XML."""
    r_pr = run._r.rPr
    if r_pr is None:
        return None

    spacing = r_pr.find(qn("w:spacing"))
    if spacing is None:
        return None

    value = spacing.get(qn("w:val"))
    return int(value) if value is not None else None


def encode_by_line_spacing(doc: Document, message: str) -> Document:
    """
    Встраивает сообщение через изменение межстрочного интервала:
    бит 0 — 12 pt, бит 1 — 14 pt.
    Один абзац хранит один бит.
    """
    bits = text_to_bits(message)
    paragraphs = [p for p in doc.paragraphs if p.text.strip()]

    if len(paragraphs) < len(bits):
        capacity_bytes = max((len(paragraphs) - 32) // 8, 0)
        raise ValueError(
            "Недостаточно абзацев для метода межстрочного интервала. "
            f"Нужно абзацев: {len(bits)}, доступно: {len(paragraphs)}. "
            f"Примерная емкость контейнера: {capacity_bytes} байт."
        )

    for paragraph, bit in zip(paragraphs, bits):
        _set_paragraph_line_spacing(
            paragraph,
            LINE_ONE_TWIPS if bit == 1 else LINE_ZERO_TWIPS,
        )

    return doc


def decode_by_line_spacing(doc: Document) -> str:
    """Извлекает сообщение, скрытое через межстрочный интервал."""
    bits = []

    for paragraph in doc.paragraphs:
        value = _get_paragraph_line_spacing(paragraph)
        if value is None:
            continue
        bits.append(1 if value >= LINE_THRESHOLD_TWIPS else 0)

    return bits_to_text(bits)


def encode_by_kerning(doc: Document, message: str) -> Document:
    """
    Встраивает сообщение через изменение кернинга/межсимвольного интервала:
    бит 0 — обычный интервал, бит 1 — расширение на 1 pt.
    Один символ хранит один бит.
    """
    bits = text_to_bits(message)
    full_text = "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text)

    if len(full_text) < len(bits):
        capacity_bytes = max((len(full_text) - 32) // 8, 0)
        raise ValueError(
            "Недостаточно символов для метода кернинга. "
            f"Нужно символов: {len(bits)}, доступно: {len(full_text)}. "
            f"Примерная емкость контейнера: {capacity_bytes} байт."
        )

    # Очищаем исходные абзацы и создаем один абзац с run-носителями.
    for paragraph in doc.paragraphs:
        paragraph.clear()

    if not doc.paragraphs:
        paragraph = doc.add_paragraph()
    else:
        paragraph = doc.paragraphs[0]

    for index, bit in enumerate(bits):
        run = paragraph.add_run(full_text[index])
        _set_run_spacing(
            run,
            KERNING_ONE_TWIPS if bit == 1 else KERNING_ZERO_TWIPS,
        )

    if len(full_text) > len(bits):
        paragraph.add_run(full_text[len(bits):])

    # Удаляем пустые абзацы, оставшиеся после очистки.
    for empty_paragraph in doc.paragraphs[1:]:
        element = empty_paragraph._element
        element.getparent().remove(element)

    return doc


def decode_by_kerning(doc: Document) -> str:
    """Извлекает сообщение, скрытое через изменение кернинга."""
    bits = []

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            value = _get_run_spacing(run)
            if value is None:
                continue
            bits.append(1 if value >= KERNING_THRESHOLD_TWIPS else 0)

    return bits_to_text(bits)


def get_capacity(doc: Document, method: str) -> int:
    """Возвращает примерную емкость контейнера в байтах без учета 32-битного заголовка."""
    if method == "line_spacing":
        carriers = len([p for p in doc.paragraphs if p.text.strip()])
    elif method == "kerning":
        carriers = len("\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text))
    else:
        raise ValueError("Неизвестный метод стеганографии.")

    return max((carriers - 32) // 8, 0)
